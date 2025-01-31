import unittest

from docx import Document


from data_ingestion.text_document import TextDocument


class TextDocumentTest(unittest.TestCase):
    def build_test_document(paragraphs):
        doc_name = "my document"

        docx_doc = Document()

        # here we build the test document
        for p in paragraphs:
            docx_doc.add_paragraph(p)


        test_filename = "test_docx_import.docx"

        docx_doc.save(test_filename)

        return TextDocument.from_docx(file_path = test_filename, name = doc_name)

    def test_that_it_builds_from_a_word_document(self):
        doc_name = "my document"
        paragraphs = [
            "[AP]: this is what was said _00:00:05_ and then this _00:01:23_"
        ]
        docx_doc = Document()

        # here we build the test document
        for p in paragraphs:
            docx_doc.add_paragraph(p)


        test_filename = "test_docx_import.docx"

        docx_doc.save(test_filename)

        doc = TextDocument.from_docx(file_path = test_filename, name = doc_name)

        self.assertEqual(doc_name, doc.name)


        for index, docx_paragraph in enumerate(docx_doc.paragraphs):
            self.assertEqual(docx_paragraph.text, doc.paragraphs[index].text)

    def test_that_it_builds_audio_labels_with_text_at_start_and_timestamp_at_end(self):
        paragraphs = [
            # "[AP]: this is what was said _00:00:05_ and then this _00:00:07_ and finally _00:00:09_"
             " this is what was said _00:00:05_ and then this _00:00:07_"
        ]
                
        test_doc = TextDocumentTest.build_test_document(paragraphs=paragraphs)

        labels = test_doc.emit_audio_labels()

        self.assertEqual(len(labels), 2)
        # TODO We need to parse speaker initials as well
        self.assert_correct_label(labels[0],' this is what was said ',0,5000)
        self.assert_correct_label(labels[1],' and then this ',5000,7000)

    def test_that_it_builds_audio_labels_with_timestamp_at_start_and_end(self):
        paragraphs = [
             "_00:00:02_ this is what was said _00:00:05_ and then this _00:00:07_"
        ]
                
        test_doc = TextDocumentTest.build_test_document(paragraphs=paragraphs)

        labels = test_doc.emit_audio_labels()

        self.assertEqual(len(labels), 2)
        # TODO We need to parse speaker initials as well
        # TODO trim whitespace
        self.assert_correct_label(labels[0],' this is what was said ',2000,5000)
        self.assert_correct_label(labels[1],' and then this ',5000,7000)

    def test_that_it_builds_audio_labels_with_text_at_start_and_end(self):
        paragraphs = [
             "But first _00:00:02_[AP] this is what was said _00:00:05_ and then this _00:00:07_ and last"
        ]
                
        test_doc = TextDocumentTest.build_test_document(paragraphs=paragraphs)

        labels = test_doc.emit_audio_labels()

        self.assertEqual(len(labels), 4)
        # TODO We need to parse speaker initials as well
        # TODO trim whitespace
        self.assert_correct_label(labels[0],'But first ',0,2000)
        self.assert_correct_label(labels[1],' this is what was said ',2000,5000)
        self.assert_correct_label(labels[2],' and then this ',5000,7000)
        # This approach to time-stamp resolution is a bit problematic
        self.assert_correct_label(labels[3],' and last',7000,7000)

    def test_that_it_builds_audio_labels_when_timestamp_at_start_and_text_at_end(self):
        paragraphs = [
             "_00:00:02_[AP] this is what was said _00:00:05_ and then this _00:00:07_ and finally!"
        ]
                
        test_doc = TextDocumentTest.build_test_document(paragraphs=paragraphs)

        labels = test_doc.emit_audio_labels()

        self.assert_correct_label(labels[0],' this is what was said ',2000,5000)
        self.assert_correct_label(labels[1],' and then this ',5000,7000)
        # What should we do in this case? Should we omit the timestamp or should we do an "averaging routine"?
        self.assert_correct_label(labels[2],' and finally!',7000,7000)

    def test_that_it_works_with_no_timestamps(self):
        paragraphs = [
            "[AP] There are no time stamps in this one"
        ]

        test_doc = TextDocumentTest.build_test_document(paragraphs=paragraphs)

        labels = test_doc.emit_audio_labels()

        self.assertEqual(len(labels),1)
    
    def test_that_it_sets_speaker_initials(self):
        paragraphs = [
             "_00:00:02_[AP] this is what was said _00:00:05_ [AP]: and then this _00:00:07_ [AP] and finally!"
        ]
                
        test_doc = TextDocumentTest.build_test_document(paragraphs=paragraphs)

        labels = test_doc.emit_audio_labels()

        for l in labels:
            self.assert_label_has_speaker_initials(l,'AP')

    def test_that_it_works_with_partial_speaker_initials(self):
        paragraphs = [
             "_00:00:02_[AP] this is what was said _00:00:05_  and then this _00:00:07_ [JB] and finally! _00:05:03"
        ]
                
        test_doc = TextDocumentTest.build_test_document(paragraphs=paragraphs)

        labels = test_doc.emit_audio_labels()

        self.assertEqual(len(labels),3)

        self.assert_label_has_speaker_initials(labels[0],'AP')

        self.assert_label_has_speaker_initials(labels[1],None)

        self.assert_label_has_speaker_initials(labels[2],'JB')

    def test_that_it_handles_multiple_paragraphs(self):
    # Note that for this case, timestamps should be carried in case there's no first time stamp
        paragraphs = [
             "_00:00:02_[AP]: this is what was said _00:00:05_  and then this _00:00:07_ [JB] and finally! _00:05:03_",
             "And this one starts the next line _00:05:55_ [BS] and keeps rolling with the time 00:07:22_"
        ]

        test_doc = TextDocumentTest.build_test_document(paragraphs=paragraphs)

        labels = test_doc.emit_audio_labels()

        self.assertEqual(len(labels),5)

        five_minutes_three_seconds_in_ms = 303000

        five_minutes_fifty_five_seconds_in_ms = 355000

        self.assert_correct_label(labels[3],'And this one starts the next line ',five_minutes_three_seconds_in_ms,five_minutes_fifty_five_seconds_in_ms)

        self.assert_label_has_speaker_initials(labels[4],'BS')
        
    
    def assert_correct_label(self,label,expected_text,expected_in_point,expected_out_point):
        self.assertEqual(label.text,expected_text) 
        self.assertEqual(label.in_point,expected_in_point)
        self.assertEqual(label.out_point,expected_out_point)

    def assert_label_has_speaker_initials(self,label,initials):
        self.assertEqual(label.speaker_initials,initials)