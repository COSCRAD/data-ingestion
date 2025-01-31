from docx import Document
import re

from data_ingestion.audio_label import AudioLabel


def parse_timestamp(text_for_timestamp):
    numeric_parts = text_for_timestamp.replace("_", "").replace("[", "").split(":")

    if len(numeric_parts) == 0:
        return 0.0

    hours, minutes, seconds = numeric_parts

    return (float(hours) * 3600 + float(minutes) * 60 + float(seconds)) * 1000  # ms


class CoscradParagraph:
    def __init__(self, text):
        self.text = text

    # The timestamp_offset allows us to shift timestamps appropriately
    def emit_audio_labels(self, first_timestamp):
        # participant_initials_delimiter_patern = r"[(a-z){2}]"

        # first_speakers_initials = re.findall(participant_initials_delimiter_patern)

        # TODO Ingect the strategy for processing labels
        time_stamp_delimiter_pattern = r"(_\d\d:\d\d:\d\d_)"
        search = re.split(time_stamp_delimiter_pattern, self.text)
        result = (
            []
            if search is None
            else list(filter(lambda s: not s.replace(" ", "") == "", search))
        )

        resolved_first_timestamp = (
            first_timestamp if first_timestamp is not None else 0.0
        )

        timestamps = []

        labels = []

        # resolve an opening time-stamp in case the passage begins with text
        if not re.match(time_stamp_delimiter_pattern, result[0]):
            timestamps.append(resolved_first_timestamp)

        for index, text in enumerate(result):
            if re.match(time_stamp_delimiter_pattern, text):
                timestamps.append(parse_timestamp(text))

        if not re.match(time_stamp_delimiter_pattern, result[-1]):
            # Do we really want to duplicate this?
            timestamps.append(timestamps[-1])

        label_index = 0

        for text in result:
            # TODO be careful about the effects of trimming white space
            if (
                not re.match(time_stamp_delimiter_pattern, text)
                and text.replace(" ", "") != ""
            ):
                search_for_initials = re.findall(r"\[[A-Za-z]{2}\]", text)

                speaker_initials_with_brackets = None

                speaker_initials = None

                if search_for_initials is not None and len(search_for_initials) == 1:
                    # TODO inject this logic
                    speaker_initials_with_brackets = search_for_initials[0]

                    # TODO be careful about trimming white space
                    text = text.replace(speaker_initials_with_brackets, "")

                    either_square_bracket_pattern = r"(\[|\])"

                    speaker_initials = re.sub(
                        either_square_bracket_pattern,
                        "",
                        speaker_initials_with_brackets,
                    )

                labels.append(
                    AudioLabel(
                        in_point=timestamps[label_index],
                        out_point=timestamps[label_index + 1],
                        text=text,
                        speaker_initials=speaker_initials,
                    )
                )

                label_index = label_index + 1

        # TODO We need to consolidate all timestamps in case there was no closing timestamp

        return labels


class TextDocument:
    def __init__(self, name):
        self.name = name

        self.paragraphs = []

    def add_paragraph(self, text):
        self.paragraphs.append(CoscradParagraph(text=text))

    def emit_audio_labels(self):
        all_labels = []

        current_timestamp = 0.0

        for p in self.paragraphs:
            for l in p.emit_audio_labels(current_timestamp):
                all_labels.append(l)

                current_timestamp = 0.0 if l.out_point is None else l.out_point

        return all_labels

    def from_docx(file_path, name):
        doc = TextDocument(name=name)

        docx_doc = Document(file_path)

        for p in docx_doc.paragraphs:
            doc.add_paragraph(p.text)

        return doc
